from __future__ import annotations

from io import BytesIO
from pathlib import Path

from fastapi import UploadFile

ALLOWED_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".json",
    ".log",
    ".pdf",
    ".docx",
}


def _decode_text(raw: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ValueError("Unable to decode file as text (tried utf-8 and gb18030)")


def _extract_pdf_text(raw: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(raw))
        parts: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
        return "\f".join(parts).strip()
    except Exception as exc:
        raise ValueError(f"Unable to parse PDF: {exc}") from exc


def _extract_docx_text(raw: bytes) -> str:
    try:
        from docx import Document

        doc = Document(BytesIO(raw))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        table_lines: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_lines.append(" | ".join(cells))
        return "\n".join([*lines, *table_lines]).strip()
    except Exception as exc:
        raise ValueError(f"Unable to parse DOCX: {exc}") from exc


def _extract_text_by_extension(*, ext: str, raw: bytes) -> str:
    if ext in {".txt", ".md", ".markdown", ".csv", ".json", ".log"}:
        return _decode_text(raw).strip()
    if ext == ".pdf":
        return _extract_pdf_text(raw)
    if ext == ".docx":
        return _extract_docx_text(raw)
    raise ValueError(f"Unsupported parser for extension: {ext}")


async def read_text_from_upload(*, upload: UploadFile, max_bytes: int) -> str:
    filename = upload.filename or ""
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext or 'unknown'}'. "
            f"Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    raw = await upload.read(max_bytes + 1)
    if not raw:
        raise ValueError("Uploaded file is empty")
    if len(raw) > max_bytes:
        raise ValueError(f"Uploaded file is too large (max {max_bytes} bytes)")

    text = _extract_text_by_extension(ext=ext, raw=raw)
    if not text:
        raise ValueError("Uploaded file contains no readable text")
    return text
